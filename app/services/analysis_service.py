import pandas as pd
import matplotlib.pyplot as plt
import os

from docx import Document
from collections import Counter
import re

def analisar_excel(caminho_arquivo, output_folder="app/static/uploads", id_arquivo=None):
    """Lê um arquivo Excel, gera métricas básicas e um gráfico, salva no static/uploads/"""
    try:
        df = pd.read_excel(caminho_arquivo)
    except Exception as e:
        return {"erro": f"Erro ao ler Excel: {e}"}

    # Coletar métricas e resumo
    resumo = {
        "linhas": len(df),
        "colunas": list(df.columns),
        "soma_colunas": df.sum(numeric_only=True).to_dict(),
        "media_colunas": df.mean(numeric_only=True).to_dict(),
        "info": df.describe(include='all').to_dict()
    }

    # Geração de gráfico (por exemplo, histograma da primeira coluna numérica)
    col_num = df.select_dtypes(include='number').columns
    grafico_path = None
    if len(col_num) > 0:
        plt.figure(figsize=(6,4))
        df[col_num[0]].hist(bins=20)
        plt.title(f"Histograma: {col_num[0]}")
        plt.xlabel(col_num[0])
        plt.ylabel('Frequência')
        # Salvar gráfico no static/uploads (com id para não sobrescrever)
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        nome_grafico = f"grafico_{id_arquivo or 'temp'}.png"
        grafico_path = os.path.join(output_folder, nome_grafico)
        plt.savefig(grafico_path)
        plt.close()
        grafico_path = grafico_path.replace("app/static/", "")  # Para usar em src no HTML

    resumo['grafico'] = grafico_path
    return resumo

def analisar_word(caminho_arquivo):
    """Lê um .docx, retorna estatísticas de palavras, parágrafos e palavras-chave"""
    try:
        doc = Document(caminho_arquivo)
        texto = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return {"erro": f"Erro ao ler Word: {e}"}

    palavras = re.findall(r'\w+', texto.lower())
    total_palavras = len(palavras)
    total_paragrafos = len(doc.paragraphs)
    palavras_mais_comuns = Counter(palavras).most_common(10)

    return {
        "total_paragrafos": total_paragrafos,
        "total_palavras": total_palavras,
        "palavras_mais_comuns": palavras_mais_comuns,
        "preview_texto": texto[:400] + ("..." if len(texto) > 400 else "")
    }